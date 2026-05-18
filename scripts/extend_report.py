from docx import Document
from docx.shared import Pt, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

BASE = os.path.dirname(__file__)
INPUT = os.path.join(BASE, '..', 'BAO_CAO_NCKH_SNIKEI_SHOP_THEO_YEU_CAU.docx')
OUTPUT = os.path.join(BASE, '..', 'BAO_CAO_NCKH_SNIKEI_SHOP_THEO_YEU_CAU_EXPANDED.docx')
IMAGES_DIR = os.path.join(BASE, '..', 'report_images')

if not os.path.exists(INPUT):
    print('Input DOCX not found:', INPUT)
    raise SystemExit(1)

print('Mở báo cáo gốc:', INPUT)
doc = Document(INPUT)

# Ensure styles: Heading 1/2/3 and Caption exist and match requirements
styles = doc.styles

def set_heading_style(name, font_name, size_pt, bold=False, italic=False):
    try:
        s = styles[name]
    except Exception:
        s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    s.font.name = font_name
    s.font.size = Pt(size_pt)
    s.font.bold = bold
    s.font.italic = italic

set_heading_style('Heading 1', 'Times New Roman', 14, bold=True)
set_heading_style('Heading 2', 'Times New Roman', 13, bold=True)
set_heading_style('Heading 3', 'Times New Roman', 13, italic=True)

# Caption style
try:
    caption = styles['Caption']
except Exception:
    caption = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
caption.font.name = 'Times New Roman'
caption.font.size = Pt(11)
caption.font.italic = True

# Normal para settings
norm = styles['Normal']
norm.font.name = 'Times New Roman'
norm.font.size = Pt(12)

# Append extended sections
print('Thêm phần Lý thuyết và Phương pháp...')
p = doc.add_paragraph('Phần 1. Lý thuyết cơ sở', style='Heading 1')
doc.add_paragraph('Trong phần này trình bày các khái niệm liên quan đến thương mại điện tử, kiến trúc MVC, mô hình dữ liệu và các kỹ thuật bảo mật cơ bản được áp dụng trong dự án Snikei Shop. Mở rộng mô tả về các thành phần, luồng dữ liệu, và ràng buộc vận hành.', style='Normal')

p = doc.add_paragraph('Phần 2. Phương pháp thực hiện', style='Heading 1')
doc.add_paragraph('Mô tả phương pháp thu thập dữ liệu, công cụ tự động hóa (Playwright, Python), phương pháp chụp ảnh màn hình, các bước kiểm thử chức năng chính (đăng nhập, thêm giỏ hàng, thanh toán) và các giả định khi thực hiện báo cáo.', style='Normal')

p = doc.add_paragraph('Phần 3. Phân tích và Kết quả', style='Heading 1')
doc.add_paragraph('Báo cáo phân tích các kết quả thu được: đánh giá giao diện, tính đúng đắn tính năng chính, thống kê nhỏ về số sản phẩm, đánh giá trải nghiệm người dùng và các vấn đề phát hiện được (ví dụ: sản phẩm không tồn tại khi dùng id mặc định).', style='Normal')

# Add a small table comparing expected vs observed
print('Thêm bảng so sánh...')
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
hdr[0].text = 'Chức năng'
hdr[1].text = 'Kỳ vọng'
hdr[2].text = 'Ghi chú (Quan sát)'
rows = [
    ('Đăng nhập', 'Người dùng admin có thể đăng nhập', 'OK'),
    ('Xem chi tiết sản phẩm', 'Hiển thị hình ảnh, giá, mô tả', 'Một số id mặc định trả về "Product not found"'),
    ('Thanh toán', 'Tạo QR/phiếu thanh toán', 'Cần thực hiện flow thêm giỏ hàng trước khi mở QR')
]
for r in rows:
    row = tbl.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

# Add detailed screenshots section with captions (reuse existing images)
print('Chèn ảnh minh họa chi tiết nếu có...')
pngs = []
if os.path.exists(IMAGES_DIR):
    pngs = sorted([f for f in os.listdir(IMAGES_DIR) if f.lower().endswith('.png')])

if pngs:
    doc.add_paragraph('Phần 4. Minh họa giao diện chi tiết', style='Heading 1')
    for i, fname in enumerate(pngs, start=1):
        path = os.path.join(IMAGES_DIR, fname)
        try:
            doc.add_picture(path, width=Cm(12))
            cap = doc.add_paragraph(f'Hình {i}. {fname.replace("_"," ").replace(".png","")}', style='Caption')
            cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception as e:
            doc.add_paragraph(f'Không chèn được ảnh {fname}: {e}', style='Normal')
else:
    doc.add_paragraph('Không tìm thấy ảnh minh họa trong thư mục report_images/', style='Normal')

# Append appendix
print('Thêm phụ lục...')
doc.add_paragraph('Phần Phụ lục', style='Heading 1')
doc.add_paragraph('A. Sơ đồ cơ sở dữ liệu: xem file database_diagram.mwb trong repository.', style='Normal')
doc.add_paragraph('B. Các script đã sử dụng để tự động hóa chụp ảnh và chèn vào báo cáo: thư mục scripts/.', style='Normal')

# Add references
print('Thêm phần Tài liệu tham khảo...')
doc.add_paragraph('Tài liệu tham khảo', style='Heading 1')
refs = [
    'Buschmann, F., et al. (1996). A system of patterns. "Pattern-Oriented Software Architecture".',
    'Gamma, E., Helm, R., Johnson, R., Vlissides, J. (1994). "Design Patterns: Elements of Reusable Object-Oriented Software".',
    'Playwright documentation. https://playwright.dev'
]
for r in refs:
    doc.add_paragraph(r, style='Normal')

# Save expanded file
print('Lưu file mở rộng:', OUTPUT)
doc.save(OUTPUT)
print('Hoàn tất.')
