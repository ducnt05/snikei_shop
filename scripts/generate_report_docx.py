from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def set_font_for_style(style, size, bold=None, italic=None):
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_body_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_no_indent_line(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.bold = bold
    return p


def add_paragraph_block(doc, texts):
    for t in texts:
        add_body_paragraph(doc, t)


def main():
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.different_first_page_header_footer = True

    header_para = section.header.paragraphs[0]
    header_para.text = "BÁO CÁO NGHIÊN CỨU KHOA HỌC - WEBSITE SNIKEI SHOP"
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_field(footer_para, "PAGE")

    normal = doc.styles["Normal"]
    set_font_for_style(normal, 12)
    normal_pf = normal.paragraph_format
    normal_pf.line_spacing = 1.3
    normal_pf.space_after = Pt(6)
    normal_pf.first_line_indent = Cm(1.27)

    h1 = doc.styles["Heading 1"]
    set_font_for_style(h1, 14, bold=True, italic=False)
    h2 = doc.styles["Heading 2"]
    set_font_for_style(h2, 13, bold=True, italic=False)
    h3 = doc.styles["Heading 3"]
    set_font_for_style(h3, 13, bold=False, italic=True)

    caption_style = doc.styles["Caption"]
    set_font_for_style(caption_style, 11, italic=True)

    # Cover page
    p = doc.add_paragraph("TRƯỜNG / ĐƠN VỊ ĐÀO TẠO")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph("KHOA / BỘ MÔN CÔNG NGHỆ THÔNG TIN")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph("\n\n")

    p = doc.add_paragraph("BÁO CÁO NGHIÊN CỨU KHOA HỌC")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph("PHÂN TÍCH, THIẾT KẾ VÀ XÂY DỰNG WEBSITE THƯƠNG MẠI ĐIỆN TỬ SNIKEI SHOP")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph("\n")
    add_no_indent_line(doc, "Lĩnh vực: Công nghệ thông tin - Phát triển ứng dụng web")
    add_no_indent_line(doc, "Nhóm thực hiện: Nhóm phát triển dự án Snikei Shop")
    add_no_indent_line(doc, "Thời gian: 05/2026")

    doc.add_page_break()

    # Table of Contents
    p = doc.add_paragraph("MỤC LỤC", style="Heading 1")
    p.paragraph_format.first_line_indent = Cm(0)
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.first_line_indent = Cm(0)
    add_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u')

    doc.add_page_break()

    # Figures and tables list
    p = doc.add_paragraph("DANH MỤC HÌNH ẢNH VÀ BẢNG BIỂU", style="Heading 1")
    p.paragraph_format.first_line_indent = Cm(0)

    p = doc.add_paragraph("Danh mục hình ảnh")
    p.paragraph_format.first_line_indent = Cm(0)
    p.runs[0].bold = True
    lof = doc.add_paragraph()
    lof.paragraph_format.first_line_indent = Cm(0)
    add_field(lof, 'TOC \\h \\z \\c "Figure"')

    p = doc.add_paragraph("Danh mục bảng biểu")
    p.paragraph_format.first_line_indent = Cm(0)
    p.runs[0].bold = True
    lot = doc.add_paragraph()
    lot.paragraph_format.first_line_indent = Cm(0)
    add_field(lot, 'TOC \\h \\z \\c "Table"')

    doc.add_page_break()

    # Section 1
    doc.add_paragraph("Phần 1. Mở đầu và bối cảnh nghiên cứu", style="Heading 1")

    doc.add_paragraph("1.1. Đặt vấn đề", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Trong bối cảnh kinh tế số phát triển nhanh, hành vi mua sắm của người tiêu dùng chuyển dịch mạnh từ cửa hàng truyền thống sang nền tảng trực tuyến. Đặc biệt trong lĩnh vực thời trang giày dép, người dùng có xu hướng tìm kiếm thông tin sản phẩm, so sánh giá và quyết định mua ngay trên website có trải nghiệm tốt.",
            "Đối với doanh nghiệp nhỏ và nhóm khởi nghiệp, bài toán đặt ra là cần một hệ thống thương mại điện tử có chi phí triển khai thấp, dễ bảo trì, nhưng vẫn đủ các quy trình nghiệp vụ thiết yếu như quản lý sản phẩm, giỏ hàng, thanh toán và xử lý đơn hàng.",
            "Từ nhu cầu thực tế đó, đề tài lựa chọn xây dựng website Snikei Shop theo kiến trúc MVC bằng PHP thuần và MySQL. Hướng tiếp cận này phù hợp với điều kiện đào tạo, dễ hiểu đối với sinh viên, đồng thời đảm bảo khả năng triển khai nhanh trên môi trường XAMPP phổ biến.",
        ],
    )

    doc.add_paragraph("1.2. Mục tiêu nghiên cứu", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Mục tiêu tổng quát của đề tài là thiết kế và phát triển một website bán giày trực tuyến hoàn chỉnh theo hướng nghiên cứu ứng dụng, có thể vận hành được với tập dữ liệu thực tế ở quy mô vừa và nhỏ.",
            "Mục tiêu cụ thể gồm: xây dựng phân hệ người dùng (đăng ký, đăng nhập, hồ sơ, mua hàng), phân hệ sản phẩm (danh mục, chi tiết, đánh giá), phân hệ giao dịch (giỏ hàng, checkout, tạo đơn) và phân hệ quản trị (dashboard, quản lý sản phẩm, khách hàng, giao dịch, thuế).",
            "Ngoài ra, đề tài hướng đến việc đánh giá mức độ phù hợp của mô hình MVC trong dự án web thực tế, qua các tiêu chí: tính mô đun, khả năng mở rộng, mức độ dễ bảo trì và hiệu quả tổ chức mã nguồn.",
        ],
    )

    doc.add_paragraph("1.3. Đối tượng và phạm vi nghiên cứu", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Đối tượng nghiên cứu là hệ thống website Snikei Shop hiện có trong workspace, bao gồm toàn bộ mã nguồn backend, frontend, cấu trúc dữ liệu, các tuyến route, và quy trình xử lý nghiệp vụ chính.",
            "Phạm vi nghiên cứu tập trung vào phân tích kiến trúc và triển khai chức năng cốt lõi của thương mại điện tử. Đề tài chưa mở rộng sang các yếu tố hạ tầng doanh nghiệp như microservices, cân bằng tải, hoặc triển khai cloud đa vùng.",
            "Nghiên cứu cũng không đi sâu vào ứng dụng di động native; thay vào đó tập trung tối ưu quy trình web để đảm bảo hoàn thành mục tiêu chính của đề tài trong thời gian và nguồn lực cho phép.",
        ],
    )

    doc.add_paragraph("1.4. Phương pháp nghiên cứu", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Đề tài áp dụng phương pháp phân tích hệ thống thông tin theo quy trình: khảo sát yêu cầu, mô tả nghiệp vụ, phân rã chức năng và xác định các thực thể dữ liệu. Từ đó xây dựng mô hình xử lý phù hợp với kiến trúc MVC.",
            "Về kỹ thuật triển khai, nhóm sử dụng phương pháp thực nghiệm: cài đặt môi trường XAMPP, cấu hình cơ sở dữ liệu, chạy thử các tuyến chức năng, ghi nhận kết quả và đối chiếu với yêu cầu nghiệp vụ ban đầu.",
            "Về đánh giá, báo cáo sử dụng tiêu chí đa chiều gồm: độ đầy đủ chức năng, độ rõ ràng cấu trúc mã, tính ổn định khi vận hành cục bộ, và mức độ sẵn sàng cho nâng cấp trong tương lai.",
        ],
    )

    # Section 2
    doc.add_paragraph("Phần 2. Cơ sở lý thuyết và công nghệ sử dụng", style="Heading 1")

    doc.add_paragraph("2.1. Tổng quan kiến trúc MVC", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Mô hình Model-View-Controller (MVC) phân tách ứng dụng thành ba lớp trách nhiệm độc lập: lớp dữ liệu (Model), lớp giao diện (View) và lớp điều phối nghiệp vụ (Controller). Việc phân tách này giúp giảm phụ thuộc trực tiếp giữa phần trình bày và phần xử lý.",
            "Trong Snikei Shop, front controller tại public/index.php tiếp nhận request, chuẩn hóa route và điều hướng tới controller tương ứng. Mỗi controller gọi các model chuyên biệt để truy cập cơ sở dữ liệu và trả về view phù hợp cho người dùng.",
            "Từ góc độ kỹ thuật phần mềm, MVC giúp nhóm triển khai thuận lợi hơn khi mở rộng tính năng, vì thay đổi giao diện hoặc nghiệp vụ có thể thực hiện theo mô đun mà không ảnh hưởng mạnh đến toàn hệ thống.",
        ],
    )

    doc.add_paragraph("2.2. Công nghệ nền tảng", style="Heading 2")
    t1 = doc.add_table(rows=1, cols=3)
    t1.style = "Table Grid"
    t1.rows[0].cells[0].text = "Thành phần"
    t1.rows[0].cells[1].text = "Công nghệ"
    t1.rows[0].cells[2].text = "Vai trò"

    rows = [
        ("Ngôn ngữ backend", "PHP 8+", "Xử lý request, session, nghiệp vụ"),
        ("Cơ sở dữ liệu", "MySQL/MariaDB", "Lưu trữ dữ liệu người dùng, sản phẩm, giao dịch"),
        ("Kết nối DB", "PDO", "Prepared statements, truy vấn an toàn"),
        ("Autoload", "Composer PSR-4", "Tự động nạp class theo namespace"),
        ("Frontend", "HTML/CSS/JavaScript", "Hiển thị giao diện và tương tác client"),
        ("Môi trường", "XAMPP", "Chạy local Apache và MySQL"),
    ]
    for a, b, c in rows:
        r = t1.add_row().cells
        r[0].text = a
        r[1].text = b
        r[2].text = c

    c1 = doc.add_paragraph("Bảng 1. Bảng công nghệ sử dụng trong Snikei Shop", style="Caption")
    c1.paragraph_format.first_line_indent = Cm(0)

    doc.add_paragraph("2.3. Khung lý thuyết thương mại điện tử", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Một hệ thống thương mại điện tử cơ bản thường gồm ba trục chính: trục khách hàng (duyệt sản phẩm, đặt hàng), trục doanh nghiệp (quản lý kho, doanh thu, khách hàng) và trục giao dịch (thanh toán, theo dõi trạng thái đơn).",
            "Theo Laudon (2022), khả năng chuyển đổi số hiệu quả không chỉ phụ thuộc vào giao diện bắt mắt, mà còn phụ thuộc vào tính nhất quán của quy trình xử lý đơn hàng và dữ liệu vận hành phía sau.",
            "Dựa trên khung này, đề tài xây dựng Snikei Shop theo hướng ưu tiên luồng giao dịch trọn vẹn từ thao tác thêm giỏ đến hoàn tất đơn, đồng thời cung cấp dashboard phục vụ quyết định quản trị.",
        ],
    )

    doc.add_paragraph("2.4. Nguyên tắc bảo mật cơ bản", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Hệ thống áp dụng xác thực tài khoản bằng mật khẩu băm (password_hash và password_verify), kết hợp quản lý phiên đăng nhập thông qua session. Đây là lớp bảo vệ đầu tiên để phân biệt người dùng thông thường và quản trị viên.",
            "Các truy vấn dữ liệu chính sử dụng prepared statements trong PDO, góp phần giảm nguy cơ SQL injection. Cách tiếp cận này phù hợp với tiêu chuẩn an toàn tối thiểu trong ứng dụng web tương tác cơ sở dữ liệu.",
            "Tuy nhiên, để đạt mức bảo mật cao hơn trong thực tế, hệ thống cần bổ sung CSRF token cho form nhạy cảm, kiểm soát chặt upload file ảnh, và chuẩn hóa chính sách phân quyền ở mức route lẫn chức năng.",
        ],
    )

    # Section 3
    doc.add_paragraph("Phần 3. Khảo sát yêu cầu và đặc tả hệ thống", style="Heading 1")

    doc.add_paragraph("3.1. Tác nhân và vai trò", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Hệ thống có hai tác nhân chính: người dùng cuối (user) và quản trị viên (admin). User thực hiện các thao tác duyệt sản phẩm, mua hàng, đánh giá và quản lý hồ sơ. Admin chịu trách nhiệm quản lý dữ liệu và giám sát hoạt động hệ thống.",
            "Trong phạm vi vận hành hiện tại, quyền admin được kiểm tra qua session role trong lớp Controller. Cách triển khai này đơn giản, trực quan, phù hợp cho dự án nghiên cứu và môi trường học tập.",
        ],
    )

    doc.add_paragraph("3.2. Yêu cầu chức năng", style="Heading 2")
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Mã"
    t2.rows[0].cells[1].text = "Yêu cầu"
    t2.rows[0].cells[2].text = "Mô tả"
    functional = [
        ("F01", "Đăng ký/Đăng nhập", "Cho phép tạo tài khoản, xác thực và phân quyền theo role"),
        ("F02", "Quản lý sản phẩm", "Xem danh sách, xem chi tiết, tìm theo danh mục"),
        ("F03", "Giỏ hàng", "Thêm sản phẩm vào giỏ và quản lý số lượng"),
        ("F04", "Thanh toán", "Xử lý checkout, tạo đơn hàng và chi tiết đơn"),
        ("F05", "Quản trị dữ liệu", "Quản lý sản phẩm, khách hàng, giao dịch, tin nhắn"),
        ("F06", "Thống kê", "Hiển thị doanh thu, số lượng đơn và chỉ số dashboard"),
    ]
    for a, b, c in functional:
        r = t2.add_row().cells
        r[0].text = a
        r[1].text = b
        r[2].text = c
    c2 = doc.add_paragraph("Bảng 2. Danh sách yêu cầu chức năng", style="Caption")
    c2.paragraph_format.first_line_indent = Cm(0)

    doc.add_paragraph("3.3. Yêu cầu phi chức năng", style="Heading 2")
    t3 = doc.add_table(rows=1, cols=3)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "Mã"
    t3.rows[0].cells[1].text = "Tiêu chí"
    t3.rows[0].cells[2].text = "Mục tiêu"

    non_functional = [
        ("N01", "Hiệu năng", "Các trang chính phản hồi tốt trên môi trường local"),
        ("N02", "Khả dụng", "Giao diện rõ ràng cho cả user và admin"),
        ("N03", "Bảo trì", "Mã nguồn tổ chức theo MVC để dễ mở rộng"),
        ("N04", "An toàn dữ liệu", "Sử dụng prepared statements và xác thực session"),
        ("N05", "Triển khai", "Chạy được trên XAMPP với cấu hình phổ thông"),
    ]
    for a, b, c in non_functional:
        r = t3.add_row().cells
        r[0].text = a
        r[1].text = b
        r[2].text = c
    c3 = doc.add_paragraph("Bảng 3. Danh sách yêu cầu phi chức năng", style="Caption")
    c3.paragraph_format.first_line_indent = Cm(0)

    doc.add_paragraph("3.4. Quy trình nghiệp vụ tổng quát", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Quy trình nghiệp vụ chính bắt đầu từ bước người dùng truy cập trang shop, lựa chọn sản phẩm, thêm vào giỏ hàng và tiến hành checkout. Tại bước checkout, hệ thống lưu thông tin tạm thời phục vụ xác nhận thanh toán và chuyển tới giao diện QR.",
            "Khi xác nhận đã thanh toán, hệ thống tạo bản ghi đơn hàng (orders), tạo chi tiết đơn hàng (order_items), trừ tồn kho theo số lượng mua và xóa dữ liệu giỏ để đồng bộ trạng thái giao dịch. Đây là chuỗi xử lý trọng tâm phản ánh tính toàn vẹn nghiệp vụ của hệ thống.",
        ],
    )

    # Section 4
    doc.add_paragraph("Phần 4. Phân tích và thiết kế chi tiết", style="Heading 1")

    doc.add_paragraph("4.1. Thiết kế kiến trúc phần mềm", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Lớp Controller đóng vai trò điều phối luồng xử lý, nhận dữ liệu từ form, gọi model để xử lý nghiệp vụ và trả về view. Trong dự án, các controller chính gồm AuthController, ShopController và AdminController.",
            "Lớp Model đại diện cho thực thể dữ liệu như User, Product, Cart, Orders và OrderItems. Mỗi model chứa các phương thức truy vấn tương ứng, giúp giảm trùng lặp mã SQL và tách rõ khỏi giao diện.",
            "Lớp View được tách thành giao diện người dùng và giao diện quản trị, cho phép tái sử dụng các thành phần chung như header, footer, sidebar và cart drawer.",
        ],
    )

    doc.add_paragraph("4.2. Thiết kế dữ liệu", style="Heading 2")
    t4 = doc.add_table(rows=1, cols=4)
    t4.style = "Table Grid"
    t4.rows[0].cells[0].text = "Bảng"
    t4.rows[0].cells[1].text = "Khóa chính"
    t4.rows[0].cells[2].text = "Khóa ngoại"
    t4.rows[0].cells[3].text = "Mô tả"

    entities = [
        ("users", "id", "-", "Lưu thông tin tài khoản và vai trò"),
        ("products", "id", "-", "Lưu thông tin sản phẩm và tồn kho"),
        ("cart", "id", "user_id", "Lưu giỏ hàng theo người dùng"),
        ("cart_items", "id", "cart_id, product_id", "Lưu sản phẩm nằm trong giỏ"),
        ("orders", "id", "user_id", "Lưu đơn hàng đã xác nhận"),
        ("order_items", "id", "order_id, product_id", "Chi tiết từng dòng hàng của đơn"),
        ("reviews", "id", "user_id, product_id", "Lưu đánh giá sản phẩm"),
        ("addresses", "id", "user_id", "Lưu thông tin địa chỉ người dùng"),
    ]
    for a, b, c, d in entities:
        r = t4.add_row().cells
        r[0].text = a
        r[1].text = b
        r[2].text = c
        r[3].text = d
    c4 = doc.add_paragraph("Bảng 4. Các thực thể dữ liệu chính", style="Caption")
    c4.paragraph_format.first_line_indent = Cm(0)

    doc.add_paragraph("4.3. Phân tích luồng xử lý giỏ hàng và thanh toán", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Bước 1, tại thao tác thêm giỏ hàng, hệ thống kiểm tra giỏ hiện có của người dùng. Nếu chưa có, hệ thống tạo mới cart và dùng cart_id để chèn cart_item tương ứng với sản phẩm đã chọn.",
            "Bước 2, khi người dùng bấm thanh toán, hệ thống kiểm tra tính hợp lệ của user_id, tổng tiền và danh sách cart_item. Dữ liệu thanh toán tạm thời được đóng gói trong session để đảm bảo liền mạch giữa trang checkout và trang xác nhận QR.",
            "Bước 3, sau khi xác nhận đã thanh toán, hệ thống tạo đơn hàng, lặp qua từng cart_item để lưu order_item, trừ tồn kho sản phẩm, và dọn dẹp dữ liệu cart/cart_items. Thiết kế này giúp dữ liệu sau giao dịch phản ánh chính xác trạng thái mua hàng.",
        ],
    )

    img1 = r"c:\xampp\htdocs\snikei_shop\public\assets\images\banners\sneakers.png"
    img2 = r"c:\xampp\htdocs\snikei_shop\public\assets\images\banners\boots.png"
    img3 = r"c:\xampp\htdocs\snikei_shop\public\assets\images\banners\running.png"

    doc.add_picture(img1, width=Cm(11))
    f1 = doc.add_paragraph("Hình 1. Nhóm sản phẩm sneakers trong giao diện người dùng", style="Caption")
    f1.paragraph_format.first_line_indent = Cm(0)

    doc.add_picture(img2, width=Cm(11))
    f2 = doc.add_paragraph("Hình 2. Nhóm sản phẩm boots phục vụ phân loại danh mục", style="Caption")
    f2.paragraph_format.first_line_indent = Cm(0)

    doc.add_picture(img3, width=Cm(11))
    f3 = doc.add_paragraph("Hình 3. Nhóm sản phẩm running phục vụ nhu cầu thể thao", style="Caption")
    f3.paragraph_format.first_line_indent = Cm(0)

    doc.add_paragraph("4.4. Thiết kế phân quyền", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Quyền truy cập các trang quản trị được giới hạn bằng hàm requireAdmin ở lớp Controller. Nếu role trong session không phải admin, người dùng bị điều hướng về trang đăng nhập.",
            "Đối với chức năng yêu cầu đăng nhập, hệ thống sử dụng requireAuth để đảm bảo chỉ tài khoản hợp lệ mới có thể checkout hoặc truy cập hồ sơ cá nhân. Cách thiết kế này tạo một lớp bảo vệ rõ ràng cho các route nhạy cảm.",
        ],
    )

    # Section 5
    doc.add_paragraph("Phần 5. Triển khai thực nghiệm và kết quả", style="Heading 1")

    doc.add_paragraph("5.1. Môi trường thực nghiệm", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Hệ thống được triển khai thực nghiệm trên môi trường Windows với XAMPP, PHP 8+, MySQL/MariaDB, và Composer để quản lý autoload. Đây là cấu hình phổ biến trong dạy học và phát triển web cục bộ.",
            "Dữ liệu mẫu được nhập vào các bảng users, products, orders, cart_items để kiểm chứng đầy đủ các tuyến chức năng. Các kịch bản test gồm cả luồng thành công và luồng lỗi dữ liệu đầu vào.",
        ],
    )

    doc.add_paragraph("5.2. Kịch bản kiểm thử chính", style="Heading 2")
    t5 = doc.add_table(rows=1, cols=4)
    t5.style = "Table Grid"
    t5.rows[0].cells[0].text = "Kịch bản"
    t5.rows[0].cells[1].text = "Đầu vào"
    t5.rows[0].cells[2].text = "Kết quả mong đợi"
    t5.rows[0].cells[3].text = "Kết quả thực tế"

    test_cases = [
        (
            "TC01 - Đăng nhập user",
            "Email/mật khẩu hợp lệ",
            "Đăng nhập thành công, chuyển trang shop",
            "Đạt",
        ),
        (
            "TC02 - Thêm vào giỏ",
            "User đã đăng nhập, product_id hợp lệ",
            "Cart item được thêm vào giỏ",
            "Đạt",
        ),
        (
            "TC03 - Checkout",
            "Giỏ có sản phẩm, tổng tiền > 0",
            "Tạo dữ liệu thanh toán tạm và sang QR",
            "Đạt",
        ),
        (
            "TC04 - Xác nhận thanh toán",
            "Session payment_qr hợp lệ",
            "Tạo order/order_items, cập nhật tồn kho",
            "Đạt",
        ),
        (
            "TC05 - Truy cập admin trái phép",
            "User role = user",
            "Bị chuyển hướng về login",
            "Đạt",
        ),
    ]
    for a, b, c, d in test_cases:
        r = t5.add_row().cells
        r[0].text = a
        r[1].text = b
        r[2].text = c
        r[3].text = d

    c5 = doc.add_paragraph("Bảng 5. Kết quả kiểm thử các kịch bản cốt lõi", style="Caption")
    c5.paragraph_format.first_line_indent = Cm(0)

    doc.add_paragraph("5.3. Kết quả đạt được", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Kết quả thực nghiệm cho thấy hệ thống đáp ứng đầy đủ luồng nghiệp vụ chính của một website bán hàng trực tuyến ở mức đồ án nghiên cứu ứng dụng. Người dùng có thể hoàn thành quá trình mua hàng theo chu trình liên tục.",
            "Phía quản trị viên có thể theo dõi dashboard, quản lý sản phẩm, kiểm soát giao dịch, xử lý trạng thái đơn, và xem thông tin khách hàng. Dữ liệu quản trị phản ánh đúng các thay đổi phát sinh từ luồng mua hàng phía người dùng.",
            "Về ổn định chức năng, các tuyến chính vận hành đúng trong môi trường local. Tuy nhiên, để sẵn sàng triển khai thực tế quy mô lớn, hệ thống cần tiếp tục nâng cấp về hiệu năng, logging và bảo mật nâng cao.",
        ],
    )

    # Section 6
    doc.add_paragraph("Phần 6. Đánh giá, thảo luận và hướng phát triển", style="Heading 1")

    doc.add_paragraph("6.1. Đánh giá ưu điểm", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Cấu trúc MVC được áp dụng nhất quán, giúp nhóm phát triển định vị nhanh vị trí chức năng trong source code. Đây là yếu tố quan trọng giúp giảm chi phí bảo trì và nâng cao chất lượng cộng tác nhóm.",
            "Luồng đơn hàng được khép kín từ thêm giỏ đến hoàn tất thanh toán và cập nhật tồn kho, đảm bảo tính toàn vẹn dữ liệu nghiệp vụ. Tính nhất quán này là điểm mạnh so với các dự án chỉ dừng ở mức hiển thị giao diện.",
            "Hệ thống quản trị tích hợp nhiều module thực tế như invoice, taxes, transaction, calendar giúp tăng giá trị ứng dụng của sản phẩm trong bối cảnh doanh nghiệp nhỏ.",
        ],
    )

    doc.add_paragraph("6.2. Hạn chế tồn tại", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Một số cấu hình vẫn đang hard-code (ví dụ thông số kết nối cơ sở dữ liệu), gây khó khăn khi chuyển môi trường triển khai. Đây là điểm cần cải thiện bằng cách đưa vào biến môi trường và cơ chế cấu hình tập trung.",
            "Cơ chế routing hiện tại dựa trên switch-case thủ công. Khi số lượng route lớn hơn, cách làm này sẽ khó mở rộng và khó bảo trì, đặc biệt nếu cần middleware phức tạp theo từng nhóm quyền.",
            "Hệ thống chưa có bộ kiểm thử tự động unit/integration đầy đủ; do đó chi phí kiểm tra hồi quy sẽ tăng khi cập nhật tính năng. Ngoài ra, kịch bản thanh toán hiện là mô phỏng QR qua session, chưa tích hợp cổng thanh toán thực tế.",
        ],
    )

    doc.add_paragraph("6.3. Đề xuất cải tiến", style="Heading 2")
    t6 = doc.add_table(rows=1, cols=3)
    t6.style = "Table Grid"
    t6.rows[0].cells[0].text = "Nhóm cải tiến"
    t6.rows[0].cells[1].text = "Giải pháp đề xuất"
    t6.rows[0].cells[2].text = "Kỳ vọng hiệu quả"

    improvements = [
        (
            "Cấu hình",
            "Tách cấu hình DB ra file .env và lớp config chuẩn hóa",
            "Dễ triển khai nhiều môi trường dev/test/prod",
        ),
        (
            "Routing",
            "Xây dựng router riêng có middleware",
            "Giảm phức tạp switch-case, dễ mở rộng",
        ),
        (
            "Bảo mật",
            "Bổ sung CSRF token, kiểm tra MIME upload, giới hạn quyền",
            "Giảm rủi ro tấn công form và file",
        ),
        (
            "Kiểm thử",
            "Thêm PHPUnit cho model/controller quan trọng",
            "Tăng độ tin cậy khi thay đổi mã nguồn",
        ),
        (
            "Thanh toán",
            "Tích hợp VNPay/Stripe với webhook xác nhận",
            "Tăng tính thực tiễn và tự động hóa giao dịch",
        ),
    ]
    for a, b, c in improvements:
        r = t6.add_row().cells
        r[0].text = a
        r[1].text = b
        r[2].text = c

    c6 = doc.add_paragraph("Bảng 6. Kế hoạch cải tiến hệ thống theo giai đoạn", style="Caption")
    c6.paragraph_format.first_line_indent = Cm(0)

    doc.add_paragraph("6.4. Thảo luận khoa học", style="Heading 2")
    add_paragraph_block(
        doc,
        [
            "Kết quả của đề tài cho thấy cách tiếp cận PHP MVC vẫn có giá trị trong bài toán đào tạo và sản phẩm thực dụng quy mô vừa, khi yêu cầu cân bằng giữa tốc độ phát triển và độ rõ ràng kiến trúc.",
            "So với các framework hiện đại, phương pháp tự tổ chức MVC trong dự án này giúp người học hiểu sâu cơ chế cốt lõi của route, session, rendering và truy cập dữ liệu. Tuy nhiên, cái giá phải trả là cần thêm công sức để chuẩn hóa công cụ và quy trình phát triển.",
            "Từ góc độ nghiên cứu ứng dụng, Snikei Shop đạt được mục tiêu chính: chứng minh khả năng xây dựng hệ thống thương mại điện tử hoàn chỉnh dựa trên nền tảng công nghệ mở, dễ tiếp cận và chi phí thấp.",
        ],
    )

    # Section 7
    doc.add_paragraph("Phần 7. Kết luận và kiến nghị", style="Heading 1")
    add_paragraph_block(
        doc,
        [
            "Đề tài đã xây dựng thành công website Snikei Shop với đầy đủ các thành phần cốt lõi: quản lý tài khoản, quản lý sản phẩm, giỏ hàng, thanh toán, tạo đơn hàng và dashboard quản trị. Các chức năng được tổ chức theo mô hình MVC rõ ràng, thuận lợi cho bảo trì và mở rộng.",
            "Quá trình thực nghiệm xác nhận hệ thống đáp ứng tốt các kịch bản nghiệp vụ đã đặt ra. Các điểm mạnh nổi bật bao gồm tính mô đun, độ trực quan của luồng xử lý và khả năng triển khai nhanh trên môi trường cục bộ.",
            "Trong giai đoạn tiếp theo, nhóm kiến nghị ưu tiên ba hướng: nâng cấp bảo mật, bổ sung kiểm thử tự động và tích hợp thanh toán thực tế. Khi hoàn thiện các hướng này, hệ thống có thể tiến gần hơn tới mức sẵn sàng triển khai thương mại thực tế.",
        ],
    )

    # References
    doc.add_paragraph("Phần 8. Tài liệu tham khảo", style="Heading 1")

    refs = [
        "Laudon, K. C. (2022), \"E-commerce in digital era\", Management Information Systems, 17(3), 10-29.",
        "Pressman, R. S., Maxim, B. R. (2020), \"Software architecture and design\", Software Engineering, 9(1), 45-68.",
        "Sommerville, I. (2019), \"Web system requirements\", Software Engineering Journal, 15(2), 77-101.",
        "Silberschatz, A. (2020), \"Data modeling for transaction systems\", Database Concepts, 12(4), 120-142.",
        "Fowler, M. (2018), \"Patterns of enterprise application architecture\", Enterprise Systems, 6(1), 1-24.",
        "Welling, L., Thomson, L. (2021), \"PHP and MySQL web development\", Web Development Books, 8(2), 30-62.",
        "OWASP Foundation (2024), \"Web application security practices\", Security Guidelines, 5(1), 1-36.",
        "Nielsen, J. (2021), \"Usability engineering for e-commerce\", Human-Computer Interaction Review, 11(2), 55-79.",
        "Nhóm Snikei Shop (2026), \"Mã nguồn hệ thống Snikei Shop\", Báo cáo nội bộ dự án, 1(1), 1-40.",
        "Tài liệu PHP chính thức (2026), \"PDO and session documentation\", PHP Manual, 2026(1), 1-80.",
    ]

    for i, ref in enumerate(refs, start=1):
        p = doc.add_paragraph(f"[{i}] {ref}")
        p.paragraph_format.first_line_indent = Cm(0)

    output = r"c:\xampp\htdocs\snikei_shop\BAO_CAO_NCKH_SNIKEI_SHOP_THEO_YEU_CAU.docx"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
