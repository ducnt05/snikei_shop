from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

BASE = os.path.dirname(__file__)
EXPANDED = os.path.join(BASE, '..', 'BAO_CAO_NCKH_SNIKEI_SHOP_THEO_YEU_CAU_EXPANDED.docx')
OUT1 = os.path.join(BASE, '..', 'BAO_CAO_NCKH_SNIKEI_SHOP_THEO_YEU_CAU_EXPANDED_WITH_IMAGES.docx')
OUT2 = os.path.join(BASE, '..', 'BAO_CAO_NCKH_SNIKEI_SHOP_THEO_YEU_CAU_WITH_IMAGES.docx')
IMAGES_DIR = os.path.join(BASE, '..', 'report_images')

if not os.path.exists(EXPANDED):
    print('Expanded docx not found:', EXPANDED)
    raise SystemExit(1)

print('Mở file expanded:', EXPANDED)
doc = Document(EXPANDED)

pngs = []
if os.path.exists(IMAGES_DIR):
    pngs = sorted([f for f in os.listdir(IMAGES_DIR) if f.lower().endswith('.png')])

if pngs:
    doc.add_paragraph('Phần minh họa giao diện (bản mở rộng)', style='Heading 1')
    for i, fname in enumerate(pngs, start=1):
        path = os.path.join(IMAGES_DIR, fname)
        try:
            doc.add_picture(path, width=Cm(12))
            cap = doc.add_paragraph(f'Hình {i}. {fname.replace("_"," ").replace(".png","")}', style='Caption')
            cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            print('Chèn:', path)
        except Exception as e:
            print('Lỗi chèn ảnh', path, e)
else:
    print('Không tìm thấy ảnh trong', IMAGES_DIR)

# Save both copies
print('Lưu:', OUT1)
doc.save(OUT1)
print('Lưu bản chung (ghi đè):', OUT2)
doc.save(OUT2)
print('Hoàn tất.')
