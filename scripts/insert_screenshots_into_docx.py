from docx import Document
from docx.shared import Cm
import os

# This script inserts screenshots from report_images/ into the report DOCX.
# Usage: python scripts/insert_screenshots_into_docx.py
# Ensure python-docx is installed and the images exist in report_images/

REPORT_DOCX = os.path.join(os.path.dirname(__file__), '..', 'BAO_CAO_NCKH_SNIKEI_SHOP_THEO_YEU_CAU.docx')
IMAGES_DIR = os.path.join(os.path.dirname(__file__), '..', 'report_images')

if not os.path.exists(REPORT_DOCX):
    print('Không tìm thấy file báo cáo:', REPORT_DOCX)
    raise SystemExit(1)

if not os.path.exists(IMAGES_DIR):
    print('Không tìm thấy thư mục ảnh:', IMAGES_DIR)
    raise SystemExit(1)

print('Mở file:', REPORT_DOCX)
doc = Document(REPORT_DOCX)

# Append a new section titled 'Minh họa giao diện' and insert all PNG images found
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

p = doc.add_paragraph('Phần minh họa giao diện', style='Heading 1')

pngs = sorted([f for f in os.listdir(IMAGES_DIR) if f.lower().endswith('.png')])
if not pngs:
    print('Không tìm thấy ảnh PNG trong', IMAGES_DIR)

for i, fname in enumerate(pngs, start=1):
    path = os.path.join(IMAGES_DIR, fname)
    caption = f'Hình {i}. {fname.replace("_"," ").replace(".png","")}'
    print('Chèn ảnh:', path)
    doc.add_picture(path, width=Cm(12))
    cap = doc.add_paragraph(caption, style='Caption')
    cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

out = REPORT_DOCX
# Save as new file to avoid overwriting accidentally
out_path = REPORT_DOCX.replace('.docx', '_WITH_IMAGES.docx')
doc.save(out_path)
print('Lưu file mới có ảnh:', out_path)
