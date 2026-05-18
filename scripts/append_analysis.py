from docx import Document
from docx.shared import Pt
import os
import subprocess

BASE = os.path.dirname(__file__)
DOCX = os.path.join(BASE, '..', 'BAO_CAO_NCKH_SNIKEI_SHOP_THEO_YEU_CAU_EXPANDED_WITH_IMAGES.docx')
if not os.path.exists(DOCX):
    print('Docx not found:', DOCX)
    raise SystemExit(1)

# Run PHP script to list products
proc = subprocess.run(['php', os.path.join(BASE, 'list_products.php')], capture_output=True, text=True)
output = proc.stdout.strip()
lines = [l for l in output.splitlines() if l.strip()]
product_count = len(lines)

print('Số sản phẩm tìm thấy:', product_count)

# Open docx and append analysis
doc = Document(DOCX)

doc.add_paragraph('Phần 5. Thống kê dữ liệu và phân tích', style='Heading 1')
doc.add_paragraph(f'Tổng số sản phẩm trong hệ thống (theo scripts/list_products.php): {product_count}', style='Normal')

# Add small list of products
doc.add_paragraph('Danh sách sản phẩm (id | tên):', style='Heading 2')
for l in lines:
    doc.add_paragraph(l, style='Normal')

# UX observations
doc.add_paragraph('Phân tích UX và đề xuất cải tiến', style='Heading 2')
doc.add_paragraph('1) Kiểm tra flow thêm sản phẩm vào giỏ: nên đảm bảo sử dụng product id hợp lệ khi tạo script tự động.', style='Normal')
doc.add_paragraph('2) Hiển thị thông báo rõ ràng khi sản phẩm không tồn tại; cải thiện giao diện 404/Not found.', style='Normal')
doc.add_paragraph('3) Thêm trạng thái xác nhận khi thêm vào giỏ để người dùng biết hành động thành công.', style='Normal')

# Save
doc.save(DOCX)
print('Đã thêm phân tích vào:', DOCX)
